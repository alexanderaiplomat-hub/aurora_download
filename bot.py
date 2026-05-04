import re
import os
import json
import glob
import time
import random
import asyncio
import logging
import subprocess
import urllib.parse
import urllib.request
from pathlib import Path
from collections import defaultdict
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, InputMediaPhoto, InputMediaVideo
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, CallbackQueryHandler, filters, ContextTypes
import yt_dlp

logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO,
    handlers=[
        logging.FileHandler('bot.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

MESSAGES = {
    'ru': {
        'downloading': '⏳ Скачиваю с {}...\nЭто может занять немного времени.',
        'downloading_progress': '⏳ Скачиваю... {}%',
        'downloading_queue': '⏳ Скачиваю с {}...\nВаш запрос в очереди, подождите.',
        'sending': '📤 Отправляю файл...',
        'downloading_subs': '⏳ Скачиваю субтитры...',
        'processing': '⏳ Обрабатываю...',
        'error_timeout': '❌ Превышено время ожидания.\n\nПричина: сервер не ответил за 120 секунд.\nПопробуй ещё раз.',
        'error_download': '❌ Ошибка при скачивании.\n\nПричина: {}\nПопробуй ещё раз.',
        'error_send': '❌ Ошибка при отправке.\n\nПричина: {}\nПопробуй ещё раз.',
        'error_not_found': '❌ Не удалось скачать.\n\nПричина: {}\nПопробуй ещё раз.',
        'error_too_large': 'Файл слишком большой для отправки (макс. 50 МБ).',
        'error_flood': 'Слишком много запросов — подожди немного.',
        'error_timed_out': 'Telegram не ответил вовремя — проблема с соединением.',
        'error_generic': '❌ Ошибка: {}',
        'error_send_generic': '❌ Ошибка при отправке: {}',
        'error_unsupported': '❌ Пожалуйста, пришли ссылку с поддерживаемой платформы.',
        'error_private': 'Это видео приватное и недоступно для скачивания.',
        'error_unavailable': 'Видео недоступно в вашем регионе или было удалено.',
        'error_instagram_auth': '❌ Instagram временно недоступен.\n\nПопробуй позже.',
        'error_tiktok_photo': 'Фото из TikTok не поддерживаются — только видео.',
        'error_stories': '📸 Сторис Instagram не поддерживаются.\n\nБот умеет скачивать Reels, посты и фото.',
        'error_playlist': 'Плейлисты не поддерживаются. Пришли ссылку на конкретный трек.',
        'error_no_rights': '❌ У вас нет прав для этой команды.',
        'subs_ready': '📝 Субтитры готовы: *{}*\n\nВыбери формат:',
        'subs_unavailable': '❌ Субтитры для этого видео недоступны.',
        'subs_empty': '❌ Субтитры пустые или недоступны.',
        'subs_expired': '❌ Субтитры устарели. Запроси их снова.',
        'subs_error_429': '❌ YouTube временно заблокировал запрос субтитров. Попробуй через несколько минут.',
        'subs_error': '❌ Ошибка при скачивании субтитров: {}',
        'desc_unavailable': 'Описание недоступно.',
        'btn_show_desc': '🔽 Название и описание',
        'btn_hide_desc': '🔼 Скрыть описание',
        'btn_show_caption': '🔽 Подпись к фото',
        'btn_hide_caption': '🔼 Скрыть подпись',
        'rate_limit': '⏱ Слишком много запросов. Подожди немного и попробуй снова.',
        'cache_cleared': '🗑 Кэш очищен. Удалено {} записей.',
        'carousel_part': 'Часть {}/{}',
        'file_large': '\n📁 Файл большой ({} МБ) — отправлен как документ',
        'link_expired': 'Ссылка устарела. Пришли её снова.',
        'connection_restore': '⏳ Соединение восстанавливается, повторяю...',
        'choose_lang': '🌍 Выбери язык / Choose language:',
        'lang_set': '✅ Язык установлен: Русский',
        'start_msg': '👋 Привет! Я бот для скачивания медиа.\n\n📥 Просто пришли мне ссылку с:\n• 🎬 YouTube — скачаю видео\n• 🎵 YouTube Music — скачаю mp3\n• 📸 Instagram (Reels, посты, фото)\n• 🎶 TikTok\n• 🐦 Twitter/X\n• И других платформ!\n\n📄 Под каждым видео есть кнопка с описанием.\n/help — все платформы\n/language — сменить язык',
        'help_msg': 'ℹ️ *Как пользоваться ботом:*\n\n1. Скопируй ссылку\n2. Отправь её мне в чат\n3. Подожди — я скачаю и пришлю файл\n\n*Поддерживаемые платформы:*\n• 🎬 YouTube и Shorts\n• 🎵 YouTube Music\n• 📸 Instagram — Reels, посты, фото\n• 🎶 TikTok\n• 🐦 Twitter/X\n• 👥 Одноклассники\n• 💙 VK Видео\n• 📺 Rutube\n• 🔄 Coub\n• 🎥 Vimeo\n• 📹 Dailymotion\n• 🎧 SoundCloud\n• 🟣 Twitch\n• 🌊 Odysee\n• 👥 Facebook\n\n*Фишки:*\n• Кнопка 🔽 под видео — название и описание\n• Повторная ссылка — мгновенно из кэша\n\n*Ограничения:*\n• Приватные видео недоступны\n• Макс. размер файла: 50 МБ\n• Фото из TikTok не поддерживаются',
    },
    'en': {
        'downloading': '⏳ Downloading from {}...\nThis may take a moment.',
        'downloading_progress': '⏳ Downloading... {}%',
        'downloading_queue': '⏳ Downloading from {}...\nYour request is queued, please wait.',
        'sending': '📤 Sending file...',
        'downloading_subs': '⏳ Downloading subtitles...',
        'processing': '⏳ Processing...',
        'error_timeout': '❌ Request timed out.\n\nReason: server did not respond in 120 seconds.\nPlease try again.',
        'error_download': '❌ Download error.\n\nReason: {}\nPlease try again.',
        'error_send': '❌ Send error.\n\nReason: {}\nPlease try again.',
        'error_not_found': '❌ Could not download.\n\nReason: {}\nPlease try again.',
        'error_too_large': 'File is too large (max 50 MB).',
        'error_flood': 'Too many requests — please wait a moment.',
        'error_timed_out': 'Telegram did not respond — connection issue.',
        'error_generic': '❌ Error: {}',
        'error_send_generic': '❌ Send error: {}',
        'error_unsupported': '❌ Please send a link from a supported platform.',
        'error_private': 'This video is private and cannot be downloaded.',
        'error_unavailable': 'Video is unavailable in your region or has been deleted.',
        'error_instagram_auth': '❌ Instagram is temporarily unavailable.\n\nPlease try again later.',
        'error_tiktok_photo': 'TikTok photos are not supported — videos only.',
        'error_stories': '📸 Instagram Stories are not supported.\n\nThe bot supports Reels, posts and photos.',
        'error_playlist': 'Playlists are not supported. Please send a link to a specific track.',
        'error_no_rights': '❌ You do not have permission to use this command.',
        'subs_ready': '📝 Subtitles ready: *{}*\n\nChoose format:',
        'subs_unavailable': '❌ Subtitles are not available for this video.',
        'subs_empty': '❌ Subtitles are empty or unavailable.',
        'subs_expired': '❌ Subtitles expired. Please request them again.',
        'subs_error_429': '❌ YouTube temporarily blocked subtitle requests. Try again in a few minutes.',
        'subs_error': '❌ Error downloading subtitles: {}',
        'desc_unavailable': 'Description unavailable.',
        'btn_show_desc': '🔽 Title and description',
        'btn_hide_desc': '🔼 Hide description',
        'btn_show_caption': '🔽 Photo caption',
        'btn_hide_caption': '🔼 Hide caption',
        'rate_limit': '⏱ Too many requests. Please wait a moment and try again.',
        'cache_cleared': '🗑 Cache cleared. {} records removed.',
        'carousel_part': 'Part {}/{}',
        'file_large': '\n📁 Large file ({} MB) — sent as document',
        'link_expired': 'Link expired. Please send it again.',
        'connection_restore': '⏳ Restoring connection, retrying...',
        'choose_lang': '🌍 Выбери язык / Choose language:',
        'lang_set': '✅ Language set: English',
        'start_msg': '👋 Hello! I am a media downloader bot.\n\n📥 Just send me a link from:\n• 🎬 YouTube — I will download the video\n• 🎵 YouTube Music — mp3\n• 📸 Instagram (Reels, posts, photos)\n• 🎶 TikTok\n• 🐦 Twitter/X\n• And other platforms!\n\n📄 Each video has a description button.\n/help — all platforms\n/language — change language',
        'help_msg': 'ℹ️ *How to use the bot:*\n\n1. Copy the link\n2. Send it to me\n3. Wait — I will download and send the file\n\n*Supported platforms:*\n• 🎬 YouTube and Shorts\n• 🎵 YouTube Music\n• 📸 Instagram — Reels, posts, photos\n• 🎶 TikTok\n• 🐦 Twitter/X\n• 👥 Odnoklassniki\n• 💙 VK Video\n• 📺 Rutube\n• 🔄 Coub\n• 🎥 Vimeo\n• 📹 Dailymotion\n• 🎧 SoundCloud\n• 🟣 Twitch\n• 🌊 Odysee\n• 👥 Facebook\n\n*Features:*\n• Button 🔽 under video — title and description\n• Repeated link — instant from cache\n\n*Limitations:*\n• Private videos unavailable\n• Max file size: 50 MB\n• TikTok photos not supported',
    },
}

def get_lang(update_or_query):
    try:
        user = None
        if hasattr(update_or_query, 'effective_user'):
            user = update_or_query.effective_user
        elif hasattr(update_or_query, 'from_user'):
            user = update_or_query.from_user
        elif hasattr(update_or_query, 'message') and update_or_query.message:
            user = update_or_query.message.from_user
        if user:
            lang = user_languages.get(user.id, 'en')
            logger.debug(f"get_lang: user_id={user.id}, lang={lang}")
            return lang
    except Exception as e:
        logger.debug(f"get_lang error: {e}")
    return 'en'

def t(update_or_query, key, *args):
    lang = get_lang(update_or_query)
    msg = MESSAGES.get(lang, MESSAGES['en']).get(key, MESSAGES['en'].get(key, key))
    if args:
        try:
            return msg.format(*args)
        except Exception:
            return msg
    return msg



def update_ytdlp():
    try:
        logger.info('Проверяю обновления yt-dlp...')
        result = subprocess.run(
            ['pip', 'install', '-U', 'yt-dlp', '--quiet'],
            capture_output=True, text=True
        )
        if 'Successfully installed' in result.stdout:
            logger.info('yt-dlp обновлён до последней версии')
        else:
            logger.info('yt-dlp актуален')
    except Exception as e:
        logger.error(f'Ошибка обновления yt-dlp: {e}')

ADMIN_ID = 1829037255  # ID администратора бота

DOWNLOAD_DIR = Path("downloads")
DOWNLOAD_DIR.mkdir(exist_ok=True)

COOKIES_FILE = Path("cookies.txt")

# Пул cookies для Instagram — бот чередует файлы
def get_instagram_cookies() -> Path | None:
    """Возвращает случайный файл cookies для Instagram"""
    pool = [Path(f) for f in glob.glob("cookies*.txt") if Path(f).exists()]
    if not pool:
        return COOKIES_FILE if COOKIES_FILE.exists() else None
    return random.choice(pool)
CACHE_FILE = Path("cache.json")

# Хранилище языков пользователей
LANGUAGES_FILE = Path("user_languages.json")

def load_languages() -> dict:
    if LANGUAGES_FILE.exists():
        try:
            with open(LANGUAGES_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return {int(k): v for k, v in data.items()}
        except Exception:
            return {}
    return {}

def save_languages(langs: dict):
    try:
        with open(LANGUAGES_FILE, 'w', encoding='utf-8') as f:
            json.dump({str(k): v for k, v in langs.items()}, f)
    except Exception as e:
        logger.error(f"Ошибка сохранения языков: {e}")

user_languages = load_languages()

processed_messages = set()

# Антиспам: chat_id -> список временных меток запросов
user_requests = defaultdict(list)
RATE_LIMIT = 10  # максимум запросов в минуту
MAX_CONCURRENT = 3  # максимум одновременных скачиваний
download_semaphore = asyncio.Semaphore(MAX_CONCURRENT)

def load_cache() -> dict:
    if CACHE_FILE.exists():
        try:
            with open(CACHE_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def save_cache(cache: dict):
    try:
        with open(CACHE_FILE, 'w', encoding='utf-8') as f:
            json.dump(cache, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Ошибка сохранения кэша: {e}")

url_cache = load_cache()

def is_supported_url(url: str) -> bool:
    patterns = [
        r'(https?://)?(www\.)?(youtube\.com|youtu\.be)/.+',
        r'(https?://)?music\.youtube\.com/.+',
        r'(https?://)?(www\.)?instagram\.com/.+',
        r'(https?://)?(www\.|vm\.|vt\.)?tiktok\.com/.+',
        r'(https?://)?(www\.)?(twitter\.com|x\.com)/.+',
        r'(https?://)?(www\.)?ok\.ru/.+',
        r'(https?://)?(www\.)?vk\.com/video.+',
        r'(https?://)?(www\.)?vkvideo\.ru/.+',
        r'(https?://)?(www\.)?rutube\.ru/.+',
        r'(https?://)?(www\.)?coub\.com/.+',
        r'(https?://)?(www\.)?vimeo\.com/.+',
        r'(https?://)?(www\.)?dailymotion\.com/.+',
        r'(https?://)?dai\.ly/.+',
        r'(https?://)?(www\.)?soundcloud\.com/.+',
        r'(https?://)?on\.soundcloud\.com/.+',
        r'(https?://)?(www\.)?twitch\.tv/.+',
        r'(https?://)?clips\.twitch\.tv/.+',
        r'(https?://)?(www\.)?odysee\.com/.+',
        r'(https?://)?(www\.|m\.)?facebook\.com/.+',
        r'(https?://)?fb\.watch/.+',
    ]
    return any(re.match(p, url) for p in patterns)

def get_platform(url: str) -> str:
    if 'music.youtube.com' in url:
        return 'YouTube Music'
    elif 'youtube.com' in url or 'youtu.be' in url:
        return 'YouTube'
    elif 'instagram.com' in url:
        return 'Instagram'
    elif 'tiktok.com' in url:
        return 'TikTok'
    elif 'twitter.com' in url or 'x.com' in url:
        return 'Twitter/X'
    elif 'ok.ru' in url:
        return 'Одноклассники'
    elif 'vk.com/video' in url or 'vkvideo.ru' in url:
        return 'VK'
    elif 'rutube.ru' in url:
        return 'Rutube'
    elif 'coub.com' in url:
        return 'Coub'
    elif 'vimeo.com' in url:
        return 'Vimeo'
    elif 'dailymotion.com' in url or 'dai.ly' in url:
        return 'Dailymotion'
    elif 'soundcloud.com' in url:
        return 'SoundCloud'
    elif 'twitch.tv' in url or 'clips.twitch.tv' in url:
        return 'Twitch'
    elif 'odysee.com' in url:
        return 'Odysee'
    elif 'facebook.com' in url or 'fb.watch' in url:
        return 'Facebook'
    return 'Unknown'

def normalize_url(url: str) -> str:
    # Для YouTube и YouTube Music сохраняем параметр v= — это ID видео/трека
    if 'youtube.com' in url or 'youtu.be' in url:
        parsed = urllib.parse.urlparse(url)
        params = urllib.parse.parse_qs(parsed.query)
        video_id = params.get('v', [''])[0]
        if video_id:
            base = f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
            return f"{base}?v={video_id}"
    return url.split('?')[0].rstrip('/')

def resolve_url(url: str) -> str:
    """Разворачивает короткие ссылки (on.soundcloud.com и др.)"""
    if 'on.soundcloud.com' in url:
        try:
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=10) as resp:
                return resp.url
        except Exception:
            pass
    return url

# Ожидающие выбор формата YouTube: chat_id -> url
pending_youtube = {}
# Хранилище скачанных субтитров: sub_id -> {txt, srt, title}
subtitles_cache = {}
_sub_counter = 0

# Хранилище описаний: id -> description text
DESCRIPTIONS_FILE = Path("descriptions.json")

def load_descriptions() -> dict:
    if DESCRIPTIONS_FILE.exists():
        try:
            with open(DESCRIPTIONS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def save_descriptions(desc: dict):
    try:
        with open(DESCRIPTIONS_FILE, 'w', encoding='utf-8') as f:
            json.dump(desc, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Ошибка сохранения описаний: {e}")

descriptions = load_descriptions()
_desc_counter = max((int(k) for k in descriptions.keys() if k.isdigit()), default=0)
# Хранилище отправленных текстов: desc_id -> text_message_id
sent_desc_messages = {}
# Счётчик скачиваний
total_downloads = 0
STATS_FILE = Path("stats.json")

def load_stats() -> dict:
    if STATS_FILE.exists():
        try:
            with open(STATS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass
    return {"total": 0}

def save_stats(stats: dict):
    try:
        with open(STATS_FILE, 'w', encoding='utf-8') as f:
            json.dump(stats, f)
    except Exception as e:
        logger.error(f"Ошибка сохранения статистики: {e}")

persistent_stats = load_stats()
if 'platforms' not in persistent_stats:
    persistent_stats['platforms'] = {}

def clean_old_cache():
    """Удаляет записи кэша старше 90 дней"""
    if not CACHE_FILE.exists():
        return
    try:
        meta_file = CACHE_FILE.with_suffix('.meta.json')
        if meta_file.exists():
            with open(meta_file, 'r', encoding='utf-8') as f:
                meta = json.load(f)
        else:
            meta = {}
        now = time.time()
        days_90 = 90 * 24 * 3600
        keys_to_delete = [k for k, v in meta.items() if now - v > days_90]
        if keys_to_delete:
            for k in keys_to_delete:
                url_cache.pop(k, None)
                meta.pop(k, None)
            save_cache(url_cache)
            with open(meta_file, 'w', encoding='utf-8') as f:
                json.dump(meta, f)
            logger.info(f'Очищено {len(keys_to_delete)} устаревших записей кэша')
    except Exception as e:
        logger.error(f'Ошибка очистки кэша: {e}')


def get_video_dimensions(filepath: Path) -> tuple[int, int]:
    """Получает реальные размеры видео через ffprobe"""
    try:
        ffprobe = r'D:\Programs\Bots\aurora_download\ffmpeg\bin\ffprobe.exe'
        result = subprocess.run([
            ffprobe, '-v', 'quiet', '-print_format', 'json',
            '-show_streams', str(filepath)
        ], capture_output=True, text=True, timeout=10)
        import json as _json
        data = _json.loads(result.stdout)
        for stream in data.get('streams', []):
            if stream.get('codec_type') == 'video':
                w = int(stream.get('width', 0))
                h = int(stream.get('height', 0))
                if w > 0 and h > 0:
                    return w, h
    except Exception:
        pass
    return 0, 0

async def download_media(url: str, chat_id: int, status_msg=None, force_audio=False, lang: str = 'en') -> tuple[Path | None, str, str, str]:
    # Возвращает: filepath, title, platform, media_type (video/audio/photo)
    url = resolve_url(url)
    platform = get_platform(url)
    if force_audio and platform == 'YouTube':
        platform = 'YouTube Music'  # используем аудио настройки

    if platform == 'YouTube Music':
        output_path = DOWNLOAD_DIR / f"{chat_id}_audio.%(ext)s"
        ydl_opts = {
            'outtmpl': str(output_path),
            'format': 'bestaudio/best',
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '192',
            }],
            'ffmpeg_location': r'D:\Programs\Bots\aurora_download\ffmpeg\bin',
            'quiet': True,
            'no_warnings': True,
            'noplaylist': True,
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            },
        }
    elif platform == 'SoundCloud':
        # Разворачиваем короткие ссылки on.soundcloud.com
        if 'on.soundcloud.com' in url:
            try:
                import requests as req_sc
                r = req_sc.get(url, allow_redirects=True, timeout=10)
                url = r.url
            except Exception:
                pass
        output_path = DOWNLOAD_DIR / f"{chat_id}_audio.%(ext)s"
        ydl_opts = {
            'outtmpl': str(output_path),
            'format': 'bestaudio/best',
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '192',
            }],
            'ffmpeg_location': r'D:\Programs\Bots\aurora_download\ffmpeg\bin',
            'quiet': True,
            'no_warnings': True,
            'noplaylist': True,
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            },
        }
    elif platform == 'Twitter/X':
        output_path = DOWNLOAD_DIR / f"{chat_id}_video.%(ext)s"
        ydl_opts = {
            'outtmpl': str(output_path),
            'format': 'best[ext=mp4]/best',
            'quiet': True,
            'no_warnings': True,
            'noplaylist': True,
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            },
        }
    elif platform == 'YouTube':
        output_path = DOWNLOAD_DIR / f"{chat_id}_video.%(ext)s"
        ydl_opts = {
            'outtmpl': str(output_path),
            'format': 'bestvideo[vcodec^=avc1][height<=720][ext=mp4]+bestaudio[ext=m4a]/bestvideo[height<=720][ext=mp4]+bestaudio[ext=m4a]/best[height<=720][ext=mp4]/best[ext=mp4]/best',
            'merge_output_format': 'mp4',
            'ffmpeg_location': r'D:\Programs\Bots\aurora_download\ffmpeg\bin',
            'quiet': True,
            'no_warnings': True,
            'noplaylist': True,
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            },
        }
    elif platform == 'Coub':
        output_path = DOWNLOAD_DIR / f"{chat_id}_%(autonumber)s.%(ext)s"
        ydl_opts = {
            'outtmpl': str(output_path),
            'format': 'bestvideo+bestaudio/best/high/med',
            'quiet': True,
            'no_warnings': True,
            'noplaylist': True,
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            },
        }
    else:
        output_path = DOWNLOAD_DIR / f"{chat_id}_%(autonumber)s.%(ext)s"
        ydl_opts = {
            'outtmpl': str(output_path),
            'format': 'best[filesize<45M]/best[filesize_approx<45M]/best',
            'quiet': True,
            'no_warnings': True,
            'noplaylist': True,
            'logtostderr': False,
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            },
        }

    # Задержка для Instagram чтобы не триггерить защиту
    if platform == 'Instagram':
        await asyncio.sleep(2)
        # Используем случайный cookies файл из пула
        instagram_cookies = get_instagram_cookies()
        if instagram_cookies:
            ydl_opts['cookiefile'] = str(instagram_cookies)
    elif COOKIES_FILE.exists():
        ydl_opts['cookiefile'] = str(COOKIES_FILE)

    try:
        for old_file in DOWNLOAD_DIR.glob(f"{chat_id}_*"):
            old_file.unlink()

        loop = asyncio.get_event_loop()
        last_progress = [0]

        def progress_hook(d):
            if d['status'] == 'downloading' and status_msg:
                total = d.get('total_bytes') or d.get('total_bytes_estimate', 0)
                downloaded = d.get('downloaded_bytes', 0)
                if total > 0:
                    percent = int(downloaded / total * 100)
                    if percent - last_progress[0] >= 20:
                        last_progress[0] = percent
                        asyncio.run_coroutine_threadsafe(
                            status_msg.edit_text(MESSAGES.get(lang, MESSAGES['en']).get('downloading_progress', '⏳ {}%').format(percent)),
                            loop
                        )

        ydl_opts['progress_hooks'] = [progress_hook]

        def _download():
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(url, download=True)
                raw_title = info.get('title', 'Media')
                description = info.get('description', '')
                artist = info.get('artist', '') or info.get('uploader', '')
                title = raw_title
                # Для YouTube Music парсим описание: "НАЗВАНИЕ · ИСПОЛНИТЕЛЬ"
                if platform == 'YouTube Music' and description:
                    for line in description.split('\n'):
                        line = line.strip()
                        if ' · ' in line and not line.startswith('Provided') and not line.startswith('Auto'):
                            parts = line.split(' · ')
                            title = parts[0].strip()
                            if len(parts) > 1:
                                artist = parts[1].strip()
                            break
                return title, artist, raw_title, description

        title, artist, orig_title, orig_desc = await loop.run_in_executor(None, _download)

        if platform == 'YouTube Music':
            display_title = f"{artist} - {title}" if artist else title
        else:
            display_title = title
            # Сохраняем описание для кнопки только для не-музыкального контента
            _desc_key = normalize_url(url)
            desc_text = orig_title
            if orig_desc and orig_desc != orig_title:
                desc_text = orig_title + "\n\n" + orig_desc
            descriptions[_desc_key] = desc_text
            save_descriptions(descriptions)

        files = sorted(DOWNLOAD_DIR.glob(f"{chat_id}_*"), key=lambda f: f.stat().st_size, reverse=True)

        if not files:
            return None, "Не удалось найти скачанный файл.", platform, 'video'

        # Для YouTube — пересобираем mp4 через ffmpeg чтобы исправить метаданные для iPhone
        if platform == 'YouTube' and len(files) == 1 and files[0].suffix.lower() == '.mp4':
            src = files[0]
            fixed = DOWNLOAD_DIR / f"{chat_id}_fixed.mp4"
            ffmpeg_path = r'D:\Programs\Bots\aurora_download\ffmpeg\bin\ffmpeg.exe'
            try:
                fix_result = subprocess.run([
                    ffmpeg_path, '-y', '-i', str(src),
                    '-c:v', 'copy', '-c:a', 'copy',
                    '-movflags', '+faststart',
                    '-map_metadata', '0',
                    str(fixed)
                ], capture_output=True, timeout=60)
                if fix_result.returncode == 0 and fixed.exists():
                    src.unlink()
                    files = [fixed]
                else:
                    if fixed.exists():
                        fixed.unlink()
            except Exception:
                if fixed.exists():
                    fixed.unlink()

        # Если файлов больше одного — это карусель
        if len(files) > 1:
            return files, display_title if platform == 'YouTube Music' else title, platform, 'carousel'

        filepath = files[0]

        # Определяем тип по расширению файла
        ext = filepath.suffix.lower()
        if ext in ('.jpg', '.jpeg', '.png', '.webp'):
            media_type = 'photo'
        elif ext in ('.mp3', '.m4a', '.ogg', '.opus', '.flac', '.wav') or platform in ('YouTube Music', 'SoundCloud'):
            media_type = 'audio'
        else:
            media_type = 'video'

        return filepath, display_title if platform == 'YouTube Music' else title, platform, media_type

    except yt_dlp.utils.DownloadError as e:
        error_msg = str(e)
        if 'Private' in error_msg or 'private' in error_msg:
            return None, 'Это видео приватное и недоступно для скачивания.', platform, 'video'
        elif 'not available' in error_msg.lower() and platform != 'Instagram':
            return None, 'Видео недоступно в вашем регионе или было удалено.', platform, 'video'
        elif platform == 'Instagram' and ('login required' in error_msg.lower() or 'rate-limit' in error_msg.lower() or 'not available' in error_msg.lower()):
            return None, 'INSTAGRAM_AUTH_ERROR', platform, 'video'
        elif 'No video formats found' in error_msg and platform == 'Instagram':
            # Пробуем скачать фото/сторис через instaloader
            try:
                import instaloader
                import re as re2
                loop2 = asyncio.get_event_loop()

                def _download_photo():
                    L = instaloader.Instaloader(
                        download_videos=True,
                        download_video_thumbnails=False,
                        download_geotags=False,
                        download_comments=False,
                        save_metadata=False,
                        post_metadata_txt_pattern='',
                        filename_pattern=f"{chat_id}_{{shortcode}}_{{mediacount}}",
                        quiet=True,
                    )
                    if COOKIES_FILE.exists():
                        try:
                            L.load_session_from_cookies(str(COOKIES_FILE))
                        except Exception:
                            pass

                    # Проверяем — сторис или обычный пост
                    stories_match = re2.search(r'/stories/([^/]+)/(\d+)', url)
                    if stories_match:
                        username = stories_match.group(1)
                        story_id = int(stories_match.group(2))
                        profile = instaloader.Profile.from_username(L.context, username)
                        for story in L.get_stories(userids=[profile.userid]):
                            for item in story.get_items():
                                if item.mediaid == story_id:
                                    L.download_storyitem(item, target=str(DOWNLOAD_DIR))
                                    return f"Story by {username}"
                        return f"Story by {username}"

                    # Извлекаем shortcode из URL
                    match = re2.search(r'/p/([^/]+)|/reel/([^/]+)', url)
                    if not match:
                        raise Exception('Не удалось извлечь shortcode из URL')
                    shortcode = match.group(1) or match.group(2)
                    post = instaloader.Post.from_shortcode(L.context, shortcode)
                    L.download_post(post, target=str(DOWNLOAD_DIR))
                    return post.caption or 'Instagram Photo'

                title2 = await loop2.run_in_executor(None, _download_photo)
                files2 = sorted(
                    [f for f in DOWNLOAD_DIR.glob(f"{chat_id}_*") if f.suffix.lower() in ('.jpg', '.jpeg', '.png', '.mp4', '.webp')],
                    key=lambda f: f.name
                )
                if not files2:
                    return None, 'Не удалось скачать фото.', platform, 'photo'
                if len(files2) > 1:
                    return files2, title2, platform, 'carousel'
                return files2[0], title2, platform, 'photo'
            except Exception as ex:
                return None, f'Не удалось скачать фото из Instagram: {str(ex)[:150]}', platform, 'photo'
        elif ('Unsupported URL' in error_msg or 'No video formats' in error_msg) and platform == 'TikTok':
            return None, 'Фото из TikTok не поддерживаются — только видео.', platform, 'photo'
        else:
            return None, f'Ошибка скачивания: {error_msg[:200]}', platform, 'video'
    except Exception as e:
        return None, f'Неожиданная ошибка: {str(e)[:200]}', platform, 'video'

async def download_subtitles(url: str, chat_id: int) -> list:
    """Скачивает субтитры и возвращает список файлов"""
    output_path = DOWNLOAD_DIR / f"{chat_id}_subs"
    ydl_opts = {
        'outtmpl': str(output_path) + '.%(ext)s',
        'writesubtitles': True,
        'writeautomaticsub': True,
        'subtitleslangs': ['ru', 'en'],
        'skip_download': True,
        'quiet': True,
        'no_warnings': True,
    }
    loop = asyncio.get_event_loop()
    def _download():
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            return info.get('title', 'Video')
    try:
        title = await loop.run_in_executor(None, _download)
    except Exception as e:
        raise Exception(str(e))
    # Находим скачанные субтитры
    srt_files = list(DOWNLOAD_DIR.glob(f"{chat_id}_subs*.vtt")) + list(DOWNLOAD_DIR.glob(f"{chat_id}_subs*.srt"))
    return srt_files, title

def vtt_to_txt(vtt_path: Path) -> str:
    """Конвертирует VTT субтитры в чистый текст"""
    text_lines = []
    with open(vtt_path, 'r', encoding='utf-8', errors='ignore') as f:
        lines = f.readlines()
    for line in lines:
        line = line.strip()
        if not line or line.startswith('WEBVTT') or line.startswith('NOTE') or '-->' in line or line.isdigit():
            continue
        # Убираем теги
        line = re.sub(r'<[^>]+>', '', line)
        if line:
            text_lines.append(line)
    # Убираем дубликаты
    seen = set()
    unique = []
    for line in text_lines:
        if line not in seen:
            seen.add(line)
            unique.append(line)
    return '\n'.join(unique)

async def prepare_subtitles(message, url: str, chat_id: int):
    """Скачивает субтитры и показывает кнопки форматов"""
    global _sub_counter
    try:
        srt_files, title = await download_subtitles(url, chat_id)
    except Exception as e:
        err = str(e)
        if '429' in err:
            await message.reply_text(t(message, "subs_error_429"))
        else:
            await message.reply_text(t(message, "subs_error", err[:150]))
        return

    if not srt_files:
        await message.reply_text(t(message, "subs_unavailable"))
        return

    # Собираем все субтитры в один текст
    all_text = ""
    all_srt = ""
    for srt_file in srt_files:
        txt = vtt_to_txt(srt_file)
        if txt.strip():
            all_text += txt + "\n"
            all_srt = srt_file.read_text(encoding='utf-8', errors='ignore')

    if not all_text.strip():
        await message.reply_text(t(message, "subs_empty"))
        for f in srt_files:
            if f.exists(): f.unlink()
        return

    # Сохраняем в хранилище
    _sub_counter += 1
    sub_id = str(_sub_counter)
    subtitles_cache[sub_id] = {
        'txt': all_text,
        'srt': all_srt,
        'title': title,
        'chat_id': chat_id,
    }

    # Чистим временные файлы
    for f in srt_files:
        if f.exists(): f.unlink()

    # Показываем кнопки форматов
    keyboard = InlineKeyboardMarkup([[
        InlineKeyboardButton("📄 TXT", callback_data=f"sub_txt_{sub_id}"),
        InlineKeyboardButton("📋 SRT", callback_data=f"sub_srt_{sub_id}"),
        InlineKeyboardButton("📝 DOCX", callback_data=f"sub_docx_{sub_id}"),
    ]])
    await message.reply_text(
        t(None, "subs_ready", title[:100]),
        reply_markup=keyboard,
        parse_mode='Markdown'
    )

async def subtitle_format_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    data = query.data

    if not data.startswith('sub_'):
        return

    parts = data.split('_', 2)
    if len(parts) < 3:
        return
    fmt = parts[1]  # txt, srt, docx
    sub_id = parts[2]

    cached = subtitles_cache.get(sub_id)
    if not cached:
        await query.message.reply_text(t(query, "subs_expired"))
        return

    title = cached['title']
    chat_id = cached['chat_id']
    safe_title = re.sub(r'[^\w\s-]', '', title)[:50].strip()

    try:
        if fmt == 'txt':
            txt_path = DOWNLOAD_DIR / f"{chat_id}_sub_out.txt"
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(f"{title}\n\n{cached['txt']}")
            with open(txt_path, 'rb') as f:
                await query.message.reply_document(
                    document=f,
                    filename=f"{safe_title}.txt",
                )
            txt_path.unlink()

        elif fmt == 'srt':
            srt_path = DOWNLOAD_DIR / f"{chat_id}_sub_out.srt"
            with open(srt_path, 'w', encoding='utf-8') as f:
                f.write(cached['srt'])
            with open(srt_path, 'rb') as f:
                await query.message.reply_document(
                    document=f,
                    filename=f"{safe_title}.srt",
                )
            srt_path.unlink()

        elif fmt == 'docx':
            try:
                from docx import Document
                doc = Document()
                doc.add_heading(title, 0)
                for line in cached['txt'].split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)
                docx_path = DOWNLOAD_DIR / f"{chat_id}_sub_out.docx"
                doc.save(str(docx_path))
                with open(docx_path, 'rb') as f:
                    await query.message.reply_document(
                        document=f,
                        filename=f"{safe_title}.docx",
                    )
                docx_path.unlink()
            except ImportError:
                await query.message.reply_text(
                    "❌ Для DOCX нужна библиотека python-docx.\n"
                    "Установи: pip install python-docx"
                )
    except Exception as e:
        await query.message.reply_text(f"❌ Ошибка: {str(e)[:150]}")

async def description_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    data = query.data

    if not data.startswith('d_'):
        return

    desc_id = data[2:]
    text = descriptions.get(desc_id)

    if not text:
        await query.message.reply_text(t(query, "desc_unavailable"))
        return

    # Проверяем — показано или скрыто
    # Определяем тип сообщения
    is_photo = bool(query.message.photo)
    is_text_msg = bool(query.message.text)  # отдельное сообщение для карусели

    btn_show = t(update, "btn_show_caption") if (is_photo or is_text_msg) else t(update, "btn_show_desc")
    btn_hide = t(query, "btn_hide_caption") if (is_photo or is_text_msg) else t(query, "btn_hide_desc")

    if desc_id in sent_desc_messages:
        # Скрываем
        keyboard = InlineKeyboardMarkup([[
            InlineKeyboardButton(btn_show, callback_data=f"d_{desc_id}")
        ]])
        if is_text_msg:
            try:
                await query.message.edit_text(".", reply_markup=keyboard)
            except Exception:
                pass
        else:
            try:
                await query.message.edit_caption(caption=None, reply_markup=keyboard)
            except Exception:
                await query.message.edit_reply_markup(reply_markup=keyboard)
        del sent_desc_messages[desc_id]
    else:
        # Показываем
        keyboard = InlineKeyboardMarkup([[
            InlineKeyboardButton(btn_hide, callback_data=f"d_{desc_id}")
        ]])
        if is_text_msg:
            try:
                await query.message.edit_text(text[:4096], reply_markup=keyboard)
            except Exception:
                pass
        else:
            try:
                await query.message.edit_caption(caption=text[:1024], reply_markup=keyboard)
            except Exception:
                await query.message.edit_reply_markup(reply_markup=keyboard)
        sent_desc_messages[desc_id] = True

async def youtube_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global _desc_counter, total_downloads, persistent_stats
    query = update.callback_query
    await query.answer()
    data = query.data
    chat_id = query.message.chat_id

    if not any(data.startswith(p) for p in ['yt_video_', 'yt_audio_', 'yt_subs_', 'yt_video_subs_', 'yt_audio_subs_']):
        return

    url = pending_youtube.get(chat_id)
    if not url:
        await query.edit_message_text(t(query, "link_expired"))
        return
    del pending_youtube[chat_id]

    if data.startswith('yt_video_subs_'):
        mode = 'video_subs'
    elif data.startswith('yt_audio_subs_'):
        mode = 'audio_subs'
    elif data.startswith('yt_video_'):
        mode = 'video'
    elif data.startswith('yt_audio_'):
        mode = 'audio'
    elif data.startswith('yt_subs_'):
        mode = 'subs'
    else:
        return

    await query.edit_message_text(t(query, "processing"))

    # Только субтитры
    if mode == 'subs':
        await query.edit_message_text(t(None, "downloading_subs"))
        await asyncio.sleep(2)
        await prepare_subtitles(query.message, url, chat_id)
        try:
            await query.message.delete()
        except Exception:
            pass
        return

    # Скачиваем видео или аудио
    force_audio = mode in ('audio', 'audio_subs')
    async with download_semaphore:
        try:
            filepath, result, plt, media_type = await asyncio.wait_for(
                download_media(url, chat_id, query.message, force_audio=force_audio),
                timeout=120
            )
        except asyncio.TimeoutError:
            await query.message.reply_text("❌ Превышено время ожидания. Попробуй ещё раз.")
            return
        except Exception as e:
            await query.message.reply_text(f"❌ Ошибка: {str(e)[:150]}")
            return

    if filepath is None:
        await query.message.reply_text(f"❌ {result}")
        return

    # Отправляем файл
    try:
        size_mb = filepath.stat().st_size / (1024 * 1024)
        with open(filepath, 'rb') as mf:
            if media_type == 'audio':
                if ' - ' in result:
                    performer, track_title = result.split(' - ', 1)
                else:
                    performer, track_title = '', result
                sent = await query.message.reply_audio(
                    audio=mf, title=track_title.strip(), performer=performer.strip(),
                    read_timeout=120, write_timeout=120,
                )
            elif size_mb <= 50:
                _desc_counter += 1
                desc_id = str(_desc_counter)
                desc_text = descriptions.get(normalize_url(url), result)
                descriptions[desc_id] = desc_text
                save_descriptions(descriptions)
                keyboard = InlineKeyboardMarkup([[
                    InlineKeyboardButton(t(update, "btn_show_desc"), callback_data=f"d_{desc_id}")
                ]])
                sent = await query.message.reply_video(
                    video=mf, reply_markup=keyboard,
                    supports_streaming=True, read_timeout=120, write_timeout=120,
                )
            else:
                sent = await query.message.reply_document(
                    document=mf, caption=f"{result[:1024]}\n📁 Файл большой ({size_mb:.1f} МБ)",
                    read_timeout=300, write_timeout=300,
                )
        await query.message.delete()
        # Субтитры дополнительно если нужно
        if mode in ('video_subs', 'audio_subs'):
            await asyncio.sleep(2)
            await prepare_subtitles(query.message, url, chat_id)
    except Exception as e:
        await query.message.reply_text(f"❌ Ошибка при отправке: {str(e)[:150]}")
    finally:
        for f in DOWNLOAD_DIR.glob(f"{chat_id}_*"):
            if f.exists():
                f.unlink()

async def language_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = InlineKeyboardMarkup([[
        InlineKeyboardButton("🇷🇺 Русский", callback_data="lang_ru"),
        InlineKeyboardButton("🇬🇧 English", callback_data="lang_en"),
    ]])
    # Показываем текущий язык пользователя
    lang = get_lang(update)
    current_lang_msg = MESSAGES[lang].get("lang_set", "")
    await update.message.reply_text(current_lang_msg, reply_markup=keyboard)

async def language_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    data = query.data
    if not data.startswith("lang_"):
        return
    if data == "lang_choose":
        keyboard = InlineKeyboardMarkup([[
            InlineKeyboardButton("🇷🇺 Русский", callback_data="lang_ru"),
            InlineKeyboardButton("🇬🇧 English", callback_data="lang_en"),
        ]])
        await query.edit_message_text(t(query, "choose_lang"), reply_markup=keyboard)
        return
    # Extract lang code - handle both lang_ru and lang_ru_start
    lang_part = data[5:]  # remove 'lang_'
    lang = lang_part.replace('_start', '')
    if lang not in MESSAGES:
        return
    user_id = query.from_user.id
    user_languages[user_id] = lang
    save_languages(user_languages)
    keyboard = InlineKeyboardMarkup([[
        InlineKeyboardButton("🇷🇺 Русский", callback_data="lang_ru"),
        InlineKeyboardButton("🇬🇧 English", callback_data="lang_en"),
    ]])
    # Определяем откуда пришёл запрос — из /start или из /language
    # callback_data содержит суффикс _start если из /start
    if data.endswith("_start"):
        # Из /start — редактируем приветственное сообщение, кнопки с _start суффиксом
        start_keyboard = InlineKeyboardMarkup([[
            InlineKeyboardButton("🇷🇺 Русский", callback_data="lang_ru_start"),
            InlineKeyboardButton("🇬🇧 English", callback_data="lang_en_start"),
        ]])
        start_msg = MESSAGES[lang].get("start_msg", "")
        await query.edit_message_text(start_msg, reply_markup=start_keyboard)
    else:
        # Из /language — короткое подтверждение
        lang_set_msg = MESSAGES[lang].get("lang_set", f"Language: {lang}")
        await query.edit_message_text(lang_set_msg, reply_markup=keyboard)

async def about_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "ℹ️ *auroradlbot*\n\n"
        "Бот для скачивания медиа из популярных социальных сетей и видеоплатформ.\n\n"
        "Просто пришли ссылку — бот скачает и пришлёт файл.\n\n"
        "Повторная ссылка отдаётся мгновенно из кэша.\n\n"
        "/help — список поддерживаемых платформ",
        parse_mode='Markdown'
    )

async def stats_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.message.chat_id != ADMIN_ID:
        await update.message.reply_text(t(update, "error_no_rights"))
        return
    cache_count = len(url_cache)
    platforms = persistent_stats.get('platforms', {})
    platforms_text = ""
    if platforms:
        sorted_platforms = sorted(platforms.items(), key=lambda x: x[1], reverse=True)
        platforms_text = "\n\n*По платформам:*\n" + "\n".join(
            f"• {p}: {c}" for p, c in sorted_platforms
        )
    await update.message.reply_text(
        f"📊 *Статистика бота:*\n\n"
        f"• В кэше: {cache_count} ссылок\n"
        f"• Скачано за сессию: {total_downloads} файлов\n"
        f"• Скачано всего: {persistent_stats.get('total', 0)} файлов"
        f"{platforms_text}",
        parse_mode='Markdown'
    )

async def clearcache_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.message.chat_id != ADMIN_ID:
        await update.message.reply_text(t(update, "error_no_rights"))
        return
    count = len(url_cache)
    url_cache.clear()
    save_cache(url_cache)
    try:
        meta_file = CACHE_FILE.with_suffix('.meta.json')
        if meta_file.exists():
            meta_file.unlink()
    except Exception:
        pass
    await update.message.reply_text(
        t(update, "cache_cleared", count)
    )

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = InlineKeyboardMarkup([[
        InlineKeyboardButton("🇷🇺 Русский", callback_data="lang_ru_start"),
        InlineKeyboardButton("🇬🇧 English", callback_data="lang_en_start"),
    ]])
    await update.message.reply_text(
        t(update, "start_msg"),
        reply_markup=keyboard
    )

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        t(update, "help_msg"),
        parse_mode='Markdown'
    )

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global _desc_counter, total_downloads, persistent_stats
    if update.edited_message:
        return

    message_id = update.message.message_id
    chat_id = update.message.chat_id

    if message_id in processed_messages:
        return

    processed_messages.add(message_id)
    # Чистим старые записи чтобы не росло бесконечно
    if len(processed_messages) > 1000:
        oldest = list(processed_messages)[:500]
        for old_id in oldest:
            processed_messages.discard(old_id)

    # Антиспам — не более 10 запросов в минуту
    now = time.time()
    user_requests[chat_id] = [t for t in user_requests[chat_id] if now - t < 60]
    if len(user_requests[chat_id]) >= RATE_LIMIT:
        try:
            await update.message.reply_text(
                t(update, "rate_limit")
            )
        except Exception:
            pass
        return
    user_requests[chat_id].append(now)

    # Ищем ссылку в тексте сообщения
    text = (update.message.text or update.message.caption or '').strip()
    url = None

    # Сначала через entities
    entities = update.message.entities or update.message.caption_entities or []
    for entity in entities:
        if entity.type == 'text_link':
            candidate = entity.url
        elif entity.type == 'url':
            try:
                text_utf16 = text.encode('utf-16-le')
                candidate = text_utf16[entity.offset*2:(entity.offset+entity.length)*2].decode('utf-16-le')
            except Exception:
                candidate = text[entity.offset:entity.offset+entity.length]
        else:
            continue
        if is_supported_url(candidate):
            url = candidate
            break

    # Если не нашли — ищем регуляркой
    if not url:
        for candidate in re.findall(r'https?://\S+', text):
            if is_supported_url(candidate):
                url = candidate
                break

    if not url:
        if update.message.chat.type == "private":
            await update.message.reply_text(
                t(update, "error_unsupported")
            )
        return

    # Проверяем сторис Instagram
    if 'instagram.com/stories/' in url:
        await update.message.reply_text(
            "📸 Сторис Instagram не поддерживаются.\n\n"
            "Бот умеет скачивать Reels, посты и фото из Instagram."
        )
        return

    # Проверяем что это не плейлист YouTube Music
    if 'music.youtube.com/playlist' in url or ('music.youtube.com' in url and 'list=' in url and 'v=' not in url):
        await update.message.reply_text(
            'Плейлисты не поддерживаются. Пожалуйста, пришли ссылку на конкретный трек.'
        )
        return

    # Проверяем кэш
    cache_key = normalize_url(url)
    if cache_key in url_cache:
        cached = url_cache[cache_key]
        file_id = cached[0]
        title = cached[1]
        file_type = cached[2]
        try:
            if file_type == 'audio':
                performer_cached = cached[3] if len(cached) > 3 else ''
                title_cached = cached[4] if len(cached) > 4 else title
                await update.message.reply_audio(
                    audio=file_id,
                    title=title_cached,
                    performer=performer_cached,
                )
            elif file_type == 'photo':
                await update.message.reply_photo(
                    photo=file_id,
                    caption=f"{title[:1024]}",
                )
            elif file_type == 'carousel':
                media_group = []
                for i, fid in enumerate(file_id):
                    caption = title[:1024] if i == 0 else None
                    # file_id для каруселей хранится как список с типами
                    fid_data = fid if isinstance(fid, str) else fid[0]
                    fid_type = 'photo' if (not isinstance(fid, str) and fid[1] == 'photo') else 'video'
                    if fid_type == 'photo':
                        media_group.append(InputMediaPhoto(media=fid_data, caption=caption))
                    else:
                        media_group.append(InputMediaVideo(media=fid_data, caption=caption))
                    if len(media_group) == 10 or i == len(file_id) - 1:
                        await update.message.reply_media_group(media=media_group, read_timeout=120, write_timeout=120)
                        media_group = []
                        if i < len(file_id) - 1:
                            await asyncio.sleep(5)
            elif file_type == 'document':
                await update.message.reply_document(
                    document=file_id,
                    caption=f"{title[:1024]}",
                )
            else:
                # Восстанавливаем кнопку описания для кэша
                _desc_counter += 1
                desc_id = str(_desc_counter)
                desc_text = title
                descriptions[desc_id] = desc_text
                save_descriptions(descriptions)
                keyboard = InlineKeyboardMarkup([[
                    InlineKeyboardButton(t(update, "btn_show_desc"), callback_data=f"d_{desc_id}")
                ]])
                await update.message.reply_video(
                    video=file_id,
                    supports_streaming=True,
                    reply_markup=keyboard,
                )
            return
        except Exception:
            del url_cache[cache_key]
            save_cache(url_cache)

    platform = get_platform(url)

    try:
        status_msg = await update.message.reply_text(
            t(update, "downloading", platform)
        )
    except Exception:
        return

    # Уведомление если бот занят
    if download_semaphore._value == 0:
        try:
            await status_msg.edit_text(
                t(update, "downloading_queue", platform)
            )
        except Exception:
            pass

    # Ограничение одновременных скачиваний
    async with download_semaphore:
        try:
            filepath, result, platform, media_type = await asyncio.wait_for(
                download_media(url, chat_id, status_msg, lang=get_lang(update)),
                timeout=120
            )
        except asyncio.TimeoutError:
            try:
                await status_msg.edit_text(
                    t(update, "error_timeout")
                )
            except Exception:
                pass
            for f in DOWNLOAD_DIR.glob(f"{chat_id}_*"):
                f.unlink()
            return
        except Exception as e:
            try:
                await status_msg.edit_text(
                    t(update, "error_download", str(e)[:150])
                )
            except Exception:
                pass
            return

    if filepath is None:
        try:
            if result == 'INSTAGRAM_AUTH_ERROR':
                # Уведомляем пользователя
                await status_msg.edit_text(
                    t(update, "error_instagram_auth")
                )
                # Уведомляем админа
                try:
                    await context.bot.send_message(
                        chat_id=ADMIN_ID,
                        text="⚠️ *Cookies для Instagram устарели!*\n\n"
                             "Нужно обновить файл `cookies.txt`.\n\n"
                             "1. Зайди в браузер где залогинен Instagram\n"
                             "2. Установи Cookie-Editor\n"
                             "3. Зайди на instagram.com\n"
                             "4. Export as Netscape → замени cookies.txt",
                        parse_mode='Markdown'
                    )
                except Exception:
                    pass
            else:
                await status_msg.edit_text(
                    t(update, "error_not_found", result)
                )
        except Exception:
            pass
        return

    try:
        await status_msg.edit_text(t(update, "sending"))

        # Если это карусель — filepath это список файлов
        if media_type == 'carousel' or isinstance(filepath, list):

            # Разбиваем на группы по 10 (лимит Telegram)
            def chunks(lst, n):
                for i in range(0, len(lst), n):
                    yield lst[i:i + n]

            groups = list(chunks(filepath, 10))

            carousel_ids = []
            for group_idx, group in enumerate(groups):
                media_group = []
                group_files = []
                for i, f in enumerate(group):
                    ext = f.suffix.lower()
                    with open(f, 'rb') as mf:
                        data = mf.read()
                    if group_idx == 0 and i == 0 and len(groups) > 1:
                        caption = t(update, "carousel_part", group_idx + 1, len(groups)) + ": " + result[:100]
                    elif group_idx == 0 and i == 0:
                        caption = result[:1024]
                    elif i == 0 and len(groups) > 1:
                        caption = t(update, "carousel_part", group_idx + 1, len(groups))
                    else:
                        caption = None

                    is_photo = ext in ('.jpg', '.jpeg', '.png', '.webp')
                    group_files.append(is_photo)
                    if is_photo:
                        media_group.append(InputMediaPhoto(media=data, caption=caption))
                    else:
                        media_group.append(InputMediaVideo(media=data, caption=caption))

                sent_msgs = await update.message.reply_media_group(media=media_group, read_timeout=120, write_timeout=120)
                # Собираем file_id из отправленных сообщений
                for msg, is_photo in zip(sent_msgs, group_files):
                    if is_photo and msg.photo:
                        carousel_ids.append([msg.photo[-1].file_id, 'photo'])
                    elif msg.video:
                        carousel_ids.append([msg.video.file_id, 'video'])

                if group_idx < len(groups) - 1:
                    await asyncio.sleep(5)

            await status_msg.delete()
            url_cache[cache_key] = (carousel_ids, result, 'carousel')
            for f in filepath:
                if f.exists():
                    f.unlink()
            save_cache(url_cache)
            return

        # Одиночное фото
        if media_type == 'photo':
            with open(filepath, 'rb') as media_file:
                # Добавляем кнопку подписи если есть текст
                if result and result.strip():
                    _desc_counter += 1
                    desc_id = str(_desc_counter)
                    descriptions[desc_id] = result
                    save_descriptions(descriptions)
                    keyboard = InlineKeyboardMarkup([[
                        InlineKeyboardButton(t(update, "btn_show_caption"), callback_data=f"d_{desc_id}")
                    ]])
                    sent = await update.message.reply_photo(
                        photo=media_file,
                        reply_markup=keyboard,
                        read_timeout=120,
                        write_timeout=120,
                    )
                else:
                    sent = await update.message.reply_photo(
                        photo=media_file,
                        read_timeout=120,
                        write_timeout=120,
                    )
            url_cache[cache_key] = (sent.photo[-1].file_id, result, 'photo')
            save_cache(url_cache)
            await status_msg.delete()
            return

        size_mb = filepath.stat().st_size / (1024 * 1024)

        with open(filepath, 'rb') as media_file:
            if platform in ('YouTube Music', 'SoundCloud'):
                logger.info(f"AUDIO SEND: result={repr(result)}, platform={platform}")
                if ' - ' in result:
                    performer, track_title = result.split(' - ', 1)
                else:
                    performer, track_title = '', result
                sent = await update.message.reply_audio(
                    audio=media_file,
                    title=track_title.strip(),
                    performer=performer.strip(),
                    read_timeout=120,
                    write_timeout=120,
                )
                url_cache[cache_key] = (sent.audio.file_id, result, 'audio', performer.strip(), track_title.strip())
            elif size_mb <= 50:
                _desc_counter += 1
                desc_id = str(_desc_counter)
                desc_text = descriptions.get(normalize_url(url), result)
                descriptions[desc_id] = desc_text
                save_descriptions(descriptions)
                keyboard = InlineKeyboardMarkup([[
                    InlineKeyboardButton(t(update, "btn_show_desc"), callback_data=f"d_{desc_id}")
                ]])
                vid_w, vid_h = get_video_dimensions(filepath)
                sent = await update.message.reply_video(
                    video=media_file,
                    reply_markup=keyboard,
                    supports_streaming=True,
                    width=vid_w if vid_w > 0 else None,
                    height=vid_h if vid_h > 0 else None,
                    read_timeout=120,
                    write_timeout=120,
                )
                url_cache[cache_key] = (sent.video.file_id, result, 'video')
            else:
                sent = await update.message.reply_document(
                    document=media_file,
                    caption=result[:1024] + t(update, "file_large", f"{size_mb:.1f}"),
                    read_timeout=300,
                    write_timeout=300,
                )
                url_cache[cache_key] = (sent.document.file_id, result, 'document')

        save_cache(url_cache)
        await status_msg.delete()
        total_downloads += 1
        persistent_stats["total"] += 1
        persistent_stats["platforms"][platform] = persistent_stats["platforms"].get(platform, 0) + 1
        save_stats(persistent_stats)

    except Exception as e:
        error_text = str(e)
        if 'Timed out' in error_text or 'TimedOut' in error_text:
            # Retry один раз при таймауте
            try:
                await asyncio.sleep(3)
                await status_msg.edit_text(t(update, "connection_restore"))
            except Exception:
                pass
            reason = t(update, "error_timed_out")
        elif 'Too Large' in error_text:
            reason = t(update, "error_too_large")
        elif 'Flood' in error_text:
            reason = t(update, "error_flood")
        else:
            reason = error_text[:150]
        try:
            await status_msg.edit_text(
                t(update, "error_send", reason)
            )
        except Exception:
            pass
    finally:
        for f in DOWNLOAD_DIR.glob(f"{chat_id}_*"):
            f.unlink()

def main():
    # Загружаем .env файл если существует
    env_file = Path(".env")
    if env_file.exists():
        with open(env_file) as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    k, v = line.split('=', 1)
                    os.environ[k.strip()] = v.strip()
    token = os.getenv("BOT_TOKEN", "8617986637:AAGCwlgwywpy6zkz-8SjoFf5WZxzaVjsAfE")

    app = ApplicationBuilder().token(token).build()

    msg_filter = filters.TEXT & ~filters.COMMAND & ~filters.UpdateType.EDITED_MESSAGE & (filters.ChatType.PRIVATE | filters.ChatType.GROUPS)

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("help", help_command))
    app.add_handler(CommandHandler("about", about_command))
    app.add_handler(CommandHandler("language", language_command))
    app.add_handler(CallbackQueryHandler(language_callback, pattern='^lang_'))
    app.add_handler(CommandHandler("stats", stats_command))
    app.add_handler(CommandHandler("clearcache", clearcache_command))
    app.add_handler(CallbackQueryHandler(youtube_callback, pattern='^yt_'))
    app.add_handler(CallbackQueryHandler(subtitle_format_callback, pattern='^sub_'))
    app.add_handler(CallbackQueryHandler(description_callback, pattern='^d_'))
    app.add_handler(MessageHandler(msg_filter, handle_message))

    # Очищаем папку downloads от мусора предыдущего запуска
    cleaned = 0
    for f in DOWNLOAD_DIR.glob("*"):
        try:
            f.unlink()
            cleaned += 1
        except Exception:
            pass
    if cleaned > 0:
        logger.info(f"Очищено {cleaned} временных файлов из папки downloads")

    update_ytdlp()
    clean_old_cache()
    logger.info("Бот запущен!")
    app.run_polling(
        drop_pending_updates=True,
        allowed_updates=["message", "edited_message", "callback_query"]
    )

if __name__ == "__main__":
    main()
